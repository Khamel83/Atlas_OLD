#!/usr/bin/env python3
import time
import os
from docx import Document

# Define login details
login_url = 'https://www.instapaper.com/user/login'
username = 'khamel83@gmail.com'
password = 'volvo87'

# Start a session
session = requests.Session()

# Create payload for login
payload = {
    'username': username,
    'password': password
}

# Perform login
login_response = session.post(login_url, data=payload)
if login_response.status_code == 200:
    print('Login successful.')
else:
    print('Login failed.')

# URL of the page with articles (please update with the correct URL if necessary)
#articles_url = 'https://www.instapaper.com/u'
articles_url = 'https://www.instapaper.com/archive'

# List to store article details
articles = []

# Number of pages to test
pages_to_test = 1  # Adjust this as needed
current_page = 0

while current_page < pages_to_test:
    # Get the page with articles
    response = session.get(articles_url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # Find all articles on the page
    article_elements = soup.find_all('article')
    print(f'Found {len(article_elements)} articles on page {current_page + 1}.')

    for article in article_elements:
        try:
            url_element = article.find('a', href=True)
            title_element = article.find('a', {'class': 'article_title'})
            summary_element = article.find('div', {'class': 'article_preview'})

            url = 'https://www.instapaper.com' + url_element['href'] if url_element else 'N/A'
            title = title_element.text.strip() if title_element else 'N/A'
            summary = summary_element.text.strip() if summary_element else 'N/A'

            articles.append({'URL': url, 'Title': title, 'Summary': summary})
        except Exception as e:
            print(f'Error parsing article: {e}')

    # Check if there's an "older articles" button to load more articles
    older_articles_button = soup.find('a', {'class': 'paginate_older'})
    if not older_articles_button:
        print('No older articles button found. Stopping.')
        break

    # Update the URL for the next set of articles
    articles_url = older_articles_button['href']
    current_page += 1

    # Wait for 1 second before loading the next page
    time.sleep(1)

# Create a Word document
doc = Document()

# Add a title for the table of contents
doc.add_heading('Table of Contents', level=1)

# Add table of contents
toc = doc.add_paragraph()
for i, article in enumerate(articles):
    toc.add_run(f"{i+1}. {article['Title']}\n")

# Add a page break
doc.add_page_break()

# Add articles to the document
for i, article in enumerate(articles):
    doc.add_heading(article['Title'], level=1)
    doc.add_paragraph(f"URL: {article['URL']}")
    doc.add_paragraph(f"Summary: {article['Summary']}")
    doc.add_paragraph("Full content:")

    for attempt in range(3):
        try:
            article_response = session.get(article['URL'])
            article_soup = BeautifulSoup(article_response.text, 'html.parser')

            # Select the main content
            main_content = article_soup.body.get_text(separator='\n')
            doc.add_paragraph(main_content)
            break
        except Exception as e:
            if attempt == 2:
                print(f"Failed to retrieve {article['URL']} after 3 attempts: {e}")
                doc.add_paragraph("[Failed to retrieve content]")
            else:
                print(f"Error retrieving {article['URL']} (attempt {attempt + 1}): {e}")
                time.sleep(2)  # Wait for 2 seconds before retrying

    # Add a page break after each article
    doc.add_page_break()

# Define the new save location
save_dir = "/Users/hr-svp-mac12/Desktop/Instapaper"
#/Users/hr-svp-mac12/Desktop/Instapaper

# Ensure the directory exists
os.makedirs(save_dir, exist_ok=True)

# Save the document
doc_path = os.path.join(save_dir, 'articles.docx')
doc.save(doc_path)

print(f'Articles saved successfully in Word document. File saved to {doc_path}.')